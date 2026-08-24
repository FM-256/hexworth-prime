#!/usr/bin/env python3
"""Rewrite the sprint documents so a student (or instructor) can type what they say and have it work.

WHY POSITIONAL, NOT TEXT-KEYED
    A text-keyed replacement corrupted Mission 3 on 2026-08-23: Missions 1 and 3 contained
    identical source lines, so Mission 1's replacement text landed in Mission 3 as well. Every
    edit here is addressed by paragraph INDEX **and** carries an expected substring that is
    verified before anything is written. If the document shifts, it refuses rather than guessing.

    EVERY edit type carries an anchor -- including insertions. An earlier version verified only
    REPLACE and INSERT_BEFORE, which left the 8-line Mission 2 SFTP block (the fix for the mission
    measured as undoable) as the single unverified splice in the patch. That is exactly the
    silent-corruption vector this file claims to have closed.

WHAT IT FIXES (all measured on a cold instance; see sprint-student-walkthrough.sh)
    * Assets: the packet referenced project files that were on NO instance. They are now baked at
      /opt/sprint-assets/ and both documents point there.
    * Mission 2 never said how the peer AUTHENTICATES; sshd had passwordauthentication=no.
    * `print(flask.__version__)` emits a DeprecationWarning that reads like an error to a student.
      It appeared TWICE in the student packet -- once in a code block, once inside a checklist
      bullet. Prose hides commands; grep the saved text, do not trust a code-block sweep.
    * "In a SECOND terminal" was impossible: access is a single noVNC console and neither tmux nor
      screen was installed. tmux is now baked in and the documents teach detaching.
    * The INSTRUCTOR RUNBOOK still carried the ORIGINAL bug untouched -- `/path/to/` placeholders,
      `apt install`, `python3 -m venv`, `pip install flask`, and `python app.py` (Ubuntu 24.04 has
      no bare `python`). It was never in scope until Chris caught it.

@catalog what    patch the sprint student packet + instructor runbook (positional, anchored)
@catalog run     python3 _tools/openstack-bridge/patch-sprint-packet.py [--check]
@catalog status  TOOL
"""
import copy
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SHARE = Path.home() / "hexworth-shared" / "openstack"

STUDENT_DOCS = [
    "OpenStack_Cloud_Security_Sprint_Student_Missions_v2.docx",
    "OpenStack_Cloud_Security_Sprint_Student_Missions_v2 (1).docx",
    "OpenStack_Cloud_Security_Sprint_Student_Missions_v2 (2).docx",
]
RUNBOOK = "OpenStack_Cloud_Security_Sprint_Instructor_Runbook_v2.docx"

FLASK_CHECK = "python3 -c \"import flask; print('flask is installed')\""

# ── student packet ───────────────────────────────────────────────────────────
STUDENT_REPLACE = [
    (19, "# 1. Start the web server:", "# Your project files are ALREADY on the instance, in /opt/sprint-assets/"),
    (20, "systemctl enable --now nginx", "sudo systemctl enable --now nginx                      # 1. start the web server"),
    (21, "sudo cp project1_index.html", "cp /opt/sprint-assets/project1_index.html ~/          # 2. copy the page to your home"),
    (22, "curl http://127.0.0.1", "curl http://127.0.0.1                                 # 5. prove it works LOCALLY first"),
    (24, "PARTNER", "# 6. Now from your PARTNER's instance (peers reach you on 'shared'):"),
    # the checklist bullet -- the copy a code-block sweep cannot see
    (72, "flask.__version__", f"☐ Confirm Flask is present: {FLASK_CHECK}"),
    (81, "flask.__version__", f"{FLASK_CHECK}  # 1. confirm flask is present"),
    # Egress was wired 2026-08-24 (router + DNS + NAT on the DevStack host), so these two lines
    # became false. Phrased as "you do not NEED to install" rather than "you cannot": everything is
    # baked in, so a class never depends on a package mirror being reachable on the day.
    (18, "have no internet",
     "# You do not need to install anything: it is all baked into ubuntu-24.04-sprint."),
    (80, "both need internet",
     "# No venv and no pip needed: flask is already installed system-wide."),
    (83, "/path/to/project3_api.py", "cp /opt/sprint-assets/project3_api.py app.py           # 3. copy the supplied API"),
    (84, "python3 app.py", "tmux new -s api                                        # 4. open a tmux session"),
    (86, "SECOND terminal", "# 6. detach from tmux with  Ctrl+b  then  d , then prove it locally:"),
    (104, "python3 project4_honeypot.py", "tmux new -s honeypot                          # 2. open a tmux session"),
    (105, "# In another terminal", "# 4. detach from tmux with:  Ctrl+b  then  d"),
    (106, "tail -f honeypot.log", "tail -f ~/honeypot.log                        # 5. watch the log fill up"),
    (108, "Authorized partner only", "# 6. On your PARTNER's instance (authorized partner only):"),
    # SSH scoping. The lab image accepts PASSWORDS (Mission 2 cannot work otherwise), so a
    # subnet-wide TCP/22 rule would let any student reach any other student's sudo account.
    # Quota is ONE instance, and a Cinder volume cannot cross projects -- so "VM #2" cannot be a
    # second simultaneous VM, and it cannot be the partner's either. The student DELETES VM #1 and
    # creates a new instance. Verified by hand 2026-08-24: the volume mounted without mkfs on the
    # new instance and proof.txt read back intact. The lesson is stronger this way: the data
    # outlives the compute entirely.
    (40, "Attach the same Cinder volume to VM #2",
     "\u2610 DELETE VM #1 (your quota is ONE instance, so VM #2 cannot exist until VM #1 is gone), "
     "create a new instance, and attach the SAME volume to it."),
    (56, "# Detach and reattach",
     "# Detach, DELETE VM #1, then reattach to a NEW instance (quota is ONE instance)"),
    (37, "Use SFTP from an approved peer",
     "\u2610 Allow TCP/22 from your PARTNER'S IP ONLY (a /32 rule), then use SFTP from that peer. "
     "Never open 22 to the whole subnet: this lab image accepts passwords."),
]
STUDENT_INSERT_BEFORE = [
    (104, "python3 project4_honeypot.py",
     ["cd ~ && cp /opt/sprint-assets/project4_honeypot.py .   # 1. copy the honeypot"]),
]
STUDENT_INSERT_AFTER = [
    (21, "sudo cp project1_index.html",
     ["nano ~/project1_index.html                             # 3. put YOUR name in it",
      "sudo cp ~/project1_index.html /var/www/html/index.html # 4. publish it"]),
    (64, "cat /srv/clouddrop/proof.txt",
     ["",
      "# SFTP DROP: your partner uploads a file to YOUR instance.",
      "# On YOUR instance, give the ubuntu account a password your partner can use:",
      "sudo passwd ubuntu",
      "",
      "# Then, on your PARTNER's instance:",
      "sftp ubuntu@<YOUR_PRIVATE_IP>",
      "# at the sftp> prompt:   put <your-file>      then:   bye"]),
    (84, "python3 app.py",
     ["python3 app.py                                         # 5. run it inside tmux"]),
    (104, "python3 project4_honeypot.py",
     ["python3 project4_honeypot.py                  # 3. run it inside tmux"]),
    (108, "Authorized partner only",
     ["cp /opt/sprint-assets/project4_generate_traffic.sh ~/",
      "chmod +x ~/project4_generate_traffic.sh"]),
    (57, "openstack server remove volume <VM1> cloud-drop",
     ["openstack server delete <VM1>          # quota is ONE instance: VM #1 must go first",
      "openstack server create <VM2> --image ubuntu-24.04-sprint --flavor ds512M --network shared"]),
]

# ── instructor runbook ───────────────────────────────────────────────────────
RUNBOOK_REPLACE = [
    (28, "sudo apt update", "# nginx, curl, nmap and ping are ALREADY BAKED INTO ubuntu-24.04-sprint."),
    (29, "apt install -y nginx", "# Nothing here needs installing. Do not spend class time on apt."),
    (30, "apt install -y curl nmap", "# The lab assets are already on the instance, in /opt/sprint-assets/"),
    (32, "sudo cp project1_index.html", "cp /opt/sprint-assets/project1_index.html ~/"),
    (94, "sudo apt update", "# flask is ALREADY INSTALLED system-wide. No venv and no pip needed."),
    (95, "apt install -y python3-venv", FLASK_CHECK),
    (96, "apt install -y curl", "mkdir -p ~/cloud-api && cd ~/cloud-api"),
    (97, "mkdir -p ~/cloud-api", "cp /opt/sprint-assets/project3_api.py app.py"),
    (98, "python3 -m venv", "tmux new -s api          # ONE console only: tmux is how you get a second shell"),
    (99, "source .venv/bin/activate", "python3 app.py           # run inside tmux, then detach: Ctrl+b then d"),
    (100, "pip install flask", "# NOTE: python3, not python -- Ubuntu 24.04 has no bare 'python' command."),
    (101, "/path/to/project3_api.py", "curl http://127.0.0.1:5000/health"),
    (102, "python app.py", "# then, from the PARTNER's instance, after allowing TCP/5000:"),
    (116, "/path/to/project4_honeypot.py", "cp /opt/sprint-assets/project4_honeypot.py ."),
    (117, "python3 project4_honeypot.py", "tmux new -s honeypot     # ONE console only: tmux gives you the second shell"),
    (119, "# Second terminal", "# Second shell: detach tmux with Ctrl+b then d, then:"),
    (124, "chmod +x project4_generate_traffic.sh", "cp /opt/sprint-assets/project4_generate_traffic.sh ."),
    # PHASE C -- the same impossible instruction that was fixed in the student packet and left
    # here. Quota is ONE instance and volumes are project-scoped, so VM #2 cannot coexist with
    # VM #1 and cannot be the partner's. Hand-verified 2026-08-24: delete VM #1, create a new
    # instance, reattach, mount WITHOUT mkfs, proof.txt intact.
    (53, "+---- detach ----> NOVA VM #2",
     "      +-- detach --> [ VM #1 DELETED ] -- attach --> NOVA VM #2 -- mount --> SAME FILES"),
    (75, "detach the Cinder volume from VM #1, attach it to VM #2",
     "Unmount the filesystem cleanly, detach the Cinder volume, then DELETE VM #1 outright. Quota is "
     "ONE instance per student and a Cinder volume cannot cross projects, so VM #2 can be neither a "
     "second simultaneous instance nor the partner's. Create a new instance, attach the same volume, "
     "mount it WITHOUT reformatting, and read proof.txt. Deleting the compute is what makes the point: "
     "the data outlives it."),
    (81, "openstack server add volume <VM2> cloud-drop",
     "openstack server delete <VM1>          # quota is ONE instance: VM #1 must go first"),
    (152, "Distribute the lab asset ZIP",
     "☐ The lab assets are BAKED INTO the image at /opt/sprint-assets/. The ZIP is a reference copy; "
     "students do not need to transfer anything."),
    (34, "partner's private IP or the 'shared' subnet",
     "Add a Neutron security-group rule for TCP/80 from the partner's private IP as a /32. Do NOT use "
     "the whole 'shared' subnet, and never 0.0.0.0/0. Do NOT associate a floating IP for this: on this "
     "cloud they are 172.24.4.0/24 and unreachable from another student's machine. Peer verification is "
     "instance-to-instance on 'shared'."),
    (73, "Allow TCP/22 only from the approved client source",
     "Allow TCP/22 from the partner's IP as a /32 ONLY. The lab image enables SSH password authentication "
     "so that Mission 2's peer SFTP can work at all, which means a subnet-wide TCP/22 rule would expose "
     "every student's instance to every other student on 'shared'. Scope it to the one partner, and take "
     "the rule back out when the mission is done."),
    (153, "Verify package repository access",
     "☐ Instances DO have internet as of 2026-08-24 (a router on shared-subnet, DNS, and NAT on the "
     "DevStack host). Do not depend on it in class: every package the missions need is baked into "
     "ubuntu-24.04-sprint. Re-run BOTH build-sprint-image.sh and wire-egress.sh after a DevStack "
     "rebuild from snapshot, because the rebuild discards both."),
]
RUNBOOK_INSERT_BEFORE = []
RUNBOOK_INSERT_AFTER = [
    (117, "python3 project4_honeypot.py", ["python3 project4_honeypot.py"]),
    (124, "chmod +x project4_generate_traffic.sh", ["chmod +x project4_generate_traffic.sh"]),
    (81, "openstack server add volume <VM2> cloud-drop",
     ["openstack server create <VM2> --image ubuntu-24.04-sprint --flavor ds512M --network shared",
      "openstack server add volume <VM2> cloud-drop"]),
]

JOBS = [(name, STUDENT_REPLACE, STUDENT_INSERT_BEFORE, STUDENT_INSERT_AFTER) for name in STUDENT_DOCS]
JOBS.append((RUNBOOK, RUNBOOK_REPLACE, RUNBOOK_INSERT_BEFORE, RUNBOOK_INSERT_AFTER))


def _retext(par: Paragraph, text: str) -> None:
    """Set a paragraph's text, keeping its first run's formatting and dropping the rest."""
    for r in par.runs[1:]:
        r._element.getparent().remove(r._element)
    if par.runs:
        par.runs[0].text = text
    else:
        par.add_run(text)


def _clone(par: Paragraph, text: str, after: bool) -> Paragraph:
    """Clone a paragraph so the copy inherits ITS style, then place it beside the original.

    Style comes from the paragraph cloned, which is why insertions anchor on the command
    paragraph itself: anchoring on the prose line above once produced a command rendered as
    body text, invisible in the code block a student reads.
    """
    new_el = copy.deepcopy(par._p)
    (par._p.addnext if after else par._p.addprevious)(new_el)
    np = Paragraph(new_el, par._parent)
    _retext(np, text)
    return np


def patch(path: Path, replace, ins_before, ins_after, check_only: bool) -> bool:
    doc = Document(str(path))
    paras = doc.paragraphs

    # Verify EVERY anchor -- replacements and both insertion kinds -- before writing anything.
    # A half-applied patch is worse than none.
    anchors = ([(i, e) for i, e, _ in replace]
               + [(i, e) for i, e, _ in ins_before]
               + [(i, e) for i, e, _ in ins_after])
    for idx, expect in anchors:
        if idx >= len(paras):
            print(f"  ✗ {path.name}: index {idx} out of range ({len(paras)} paragraphs)")
            return False
        if expect not in paras[idx].text:
            print(f"  ✗ {path.name}: anchor mismatch at [{idx}], expected {expect!r}, "
                  f"found {paras[idx].text[:60]!r}")
            return False
    if check_only:
        print(f"  ✓ {path.name}: all {len(anchors)} anchors match")
        return True

    for idx, _, new in replace:
        _retext(paras[idx], new)

    # Insertions bottom-up across BOTH lists together: a higher index handled first cannot
    # disturb a lower one. Separate passes per list would corrupt the ordering.
    jobs = ([(i, True, lines) for i, _, lines in ins_after]
            + [(i, False, lines) for i, _, lines in ins_before])
    for idx, after, lines in sorted(jobs, key=lambda t: -t[0]):
        for line in (reversed(lines) if after else lines):
            _clone(paras[idx], line, after)

    doc.save(str(path))
    n = sum(len(l) for _, _, l in ins_after) + sum(len(l) for _, _, l in ins_before)
    print(f"  ✓ {path.name}: {len(replace)} replaced, {n} inserted")
    return True


if __name__ == "__main__":
    check = "--check" in sys.argv
    ok = all(patch(SHARE / name, r, b, a, check) for name, r, b, a in JOBS)
    sys.exit(0 if ok else 1)
